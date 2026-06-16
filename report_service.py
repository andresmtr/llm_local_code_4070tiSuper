import io
import json
import os
import logging
from datetime import datetime

import requests
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from dotenv import load_dotenv

from data_fetching_service import get_records, get_municipio_coords, resolve_record, YEAR_FIELD

load_dotenv()
logger = logging.getLogger(__name__)

OPENROUTER_API_KEY = os.getenv('OPENROUTER_API_KEY', '')
OPENROUTER_MODEL = os.getenv('OPENROUTER_MODEL', 'nvidia/nemotron-3-ultra-550b-a55b:free')

MONTH_ORDER = ['enero','febrero','marzo','abril','mayo','junio','julio','agosto','septiembre','octubre','noviembre','diciembre']


def _parse_year(val):
    if not val:
        return None
    if isinstance(val, (int, float)):
        return int(val)
    try:
        return int(str(val).strip()[:4])
    except (ValueError, TypeError):
        return None


def _clean_context(name):
    prefixes = ['1 Lesiones no Fatales por ', '2 Lesiones no Fatales por ',
                '3 Lesiones no Fatales contra Niños, Niñas y Adolescentes por ',
                '4 Lesiones no Fatales por ', '5 Lesiones no Fatales por ',
                '6 Lesiones no Fatales por ', '7 Lesiones no Fatales por ',
                '8 Lesiones no Fatales por ', '9 Lesiones no Fatales por ']
    for p in prefixes:
        if name.startswith(p):
            return name[len(p):]
    return name


def _agg(records, key, limit=None):
    d = {}
    for r in records:
        v = r.get(key) or 'Sin info'
        d[v] = d.get(v, 0) + 1
    items = sorted(d.items(), key=lambda x: -x[1])
    if limit:
        items = items[:limit]
    return items


def _ai_enhance(prompt, system_prompt=None):
    if not OPENROUTER_API_KEY:
        return None
    try:
        messages = []
        if system_prompt:
            messages.append({'role': 'system', 'content': system_prompt})
        messages.append({'role': 'user', 'content': prompt})
        resp = requests.post(
            'https://openrouter.ai/api/v1/chat/completions',
            headers={
                'Authorization': f'Bearer {OPENROUTER_API_KEY}',
                'Content-Type': 'application/json',
            },
            json={
                'model': OPENROUTER_MODEL,
                'messages': messages,
                'max_tokens': 600,
                'temperature': 0.7,
                'reasoning': {'enabled': True},
            },
            timeout=120,
        )
        resp.raise_for_status()
        data = resp.json()
        return data['choices'][0]['message']['content']
    except Exception as e:
        logger.warning('OpenRouter API error: %s', e)
        return None


def _set_cell(cell, text, bold=False, size=10, color=None):
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        _set_cell(table.rows[0].cells[i], h, bold=True, size=9)
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            _set_cell(row.cells[i], str(val), size=9)
    return table


def generate_report(departamento='', municipio='', year=''):
    records = get_records()
    geo_lookup = get_municipio_coords()
    if departamento or municipio or year:
        filtered = []
        for r in records:
            if departamento or municipio:
                resolved = resolve_record(r, geo_lookup)
                if departamento and resolved['dept_name'] != departamento:
                    continue
                if municipio and resolved['mun_name'] != municipio:
                    continue
            if year:
                y = _parse_year(r.get(YEAR_FIELD))
                if y is None or str(y) != year:
                    continue
            filtered.append(r)
        records = filtered

    total = len(records)

    by_gender = dict(_agg(records, 'identidad_de_genero'))
    by_age = dict(_agg(records, 'grupo_de_edad_de_la_victima'))
    by_context_raw = _agg(records, 'contexto_de_violencia', 10)
    by_context = [(_clean_context(k), v) for k, v in by_context_raw]
    by_scenario = dict(_agg(records, 'escenario_del_hecho', 10))
    by_month_raw = _agg(records, 'mes_del_hecho')
    by_month_raw.sort(key=lambda x: MONTH_ORDER.index(x[0]) if x[0] in MONTH_ORDER else 99)
    by_causal = dict(_agg(records, 'mecanismo_causal', 10))
    by_dept = dict(_agg(records, 'departamento_del_hecho_dane'))
    by_injury = dict(_agg(records, 'dias_de_incapacidad_medicolegal'))
    by_zona = dict(_agg(records, 'zona_del_hecho'))

    doc = Document()

    # Estilos base
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    # --- PORTADA ---
    for _ in range(6):
        doc.add_paragraph('')
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('INFORME DE LESIONES\nNO FATALES DE CAUSA EXTERNA')
    run.font.size = Pt(26)
    run.bold = True
    run.font.color.rgb = RGBColor(13, 27, 74)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f'\nDatos preliminares {datetime.now().strftime("%B %Y")}')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph('')
    filtro_text = f'Departamento: {departamento or "Todos"} | Municipio: {municipio or "Todos"} | Año: {year or "Todos"}'
    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run(filtro_text)
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_paragraph('')
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tp.add_run(f'Total de registros: {total:,}')
    run.font.size = Pt(16)
    run.bold = True
    run.font.color.rgb = RGBColor(13, 27, 74)

    doc.add_page_break()

    # --- RESUMEN EJECUTIVO ---
    h = doc.add_heading('1. Resumen Ejecutivo', level=1)
    resumen_datos = (
        f"Se analizaron {total} registros de lesiones no fatales de causa externa "
        f"en Colombia. Los datos incluyen información sobre víctimas, agresores, "
        f"contextos de violencia, ubicaciones geográficas y más. "
        f"Distribución por género: {dict(by_gender)}. "
        f"Principales contextos de violencia: {dict(by_context)}. "
        f"Distribución por zona: {dict(by_zona)}."
    )

    ai_text = _ai_enhance(
        f"Genera un resumen ejecutivo conciso (máximo 3 párrafos) para un informe "
        f"de lesiones no fatales de causa externa en Colombia con estos datos:\n{resumen_datos}\n\n"
        f"El resumen debe ser profesional, objetivo y destacar los hallazgos principales.",
        system_prompt="Eres un analista de datos forenses que redacta informes profesionales en español."
    )

    if ai_text:
        p = doc.add_paragraph(ai_text)
    else:
        p = doc.add_paragraph(
            f"Este informe presenta el análisis de {total} registros de lesiones no fatales "
            f"de causa externa. Se examinan variables como género, grupo de edad, contexto de "
            f"violencia, ubicación geográfica, mecanismo causal y días de incapacidad, "
            f"proporcionando una visión integral de la problemática."
        )

    doc.add_page_break()

    # --- DATOS GENERALES ---
    h = doc.add_heading('2. Datos Generales', level=1)

    doc.add_heading('2.1 Distribución por Género', level=2)
    gender_data = [[g, str(c), f'{c/total*100:.1f}%'] for g, c in sorted(by_gender.items(), key=lambda x: -x[1])]
    _add_table(doc, ['Género', 'Cantidad', '%'], gender_data)

    doc.add_paragraph('')
    doc.add_heading('2.2 Distribución por Grupo de Edad', level=2)
    age_data = [[a, str(c), f'{c/total*100:.1f}%'] for a, c in sorted(by_age.items(), key=lambda x: -x[1])]
    _add_table(doc, ['Grupo de Edad', 'Cantidad', '%'], age_data)

    doc.add_paragraph('')
    doc.add_heading('2.3 Distribución por Zona', level=2)
    zona_data = [[z, str(c), f'{c/total*100:.1f}%'] for z, c in sorted(by_zona.items(), key=lambda x: -x[1])]
    _add_table(doc, ['Zona', 'Cantidad', '%'], zona_data)

    doc.add_page_break()

    # --- CONTEXTO DE VIOLENCIA ---
    h = doc.add_heading('3. Análisis de Contexto de Violencia', level=1)
    ctx_data = [[_clean_context(c) if len(_clean_context(c)) > 3 else c, str(v), f'{v/total*100:.1f}%'] for c, v in by_context_raw]
    _add_table(doc, ['Contexto de Violencia', 'Cantidad', '%'], ctx_data)

    doc.add_paragraph('')
    ai_context = _ai_enhance(
        f"Analiza los siguientes datos de contextos de violencia en lesiones no fatales:\n{dict(by_context)}\n\n"
        f"Genera un análisis de 2-3 párrafos explicando los patrones observados, "
        f"los contextos más frecuentes y sus implicaciones.",
        system_prompt="Eres un criminólogo experto en análisis de violencia en Colombia."
    )
    if ai_context:
        doc.add_paragraph(ai_context)
    else:
        top_context = by_context[0][0] if by_context else 'N/A'
        doc.add_paragraph(
            f"El contexto de violencia más frecuente es '{top_context}', representando "
            f"el {by_context[0][1]/total*100:.1f}% de los casos registrados."
        )

    doc.add_page_break()

    # --- UBICACIÓN GEOGRÁFICA ---
    h = doc.add_heading('4. Distribución Geográfica', level=1)

    doc.add_heading('4.1 Por Departamento', level=2)
    dept_data = [[d, str(c), f'{c/total*100:.1f}%'] for d, c in sorted(by_dept.items(), key=lambda x: -x[1])]
    _add_table(doc, ['Departamento', 'Cantidad', '%'], dept_data)

    doc.add_paragraph('')
    doc.add_heading('4.2 Tendencia Mensual', level=2)
    month_data = [[m.capitalize(), str(by_month_raw[i][1])] for i, (m, _) in enumerate(by_month_raw)]
    _add_table(doc, ['Mes', 'Cantidad'], month_data)

    doc.add_page_break()

    # --- MECANISMOS Y ESCENARIOS ---
    h = doc.add_heading('5. Mecanismos y Escenarios', level=1)

    doc.add_heading('5.1 Mecanismo Causal', level=2)
    causal_data = [[m, str(c)] for m, c in sorted(by_causal.items(), key=lambda x: -x[1])]
    _add_table(doc, ['Mecanismo Causal', 'Cantidad'], causal_data)

    doc.add_paragraph('')
    doc.add_heading('5.2 Escenario del Hecho', level=2)
    scen_data = [[s, str(c)] for s, c in sorted(by_scenario.items(), key=lambda x: -x[1])]
    _add_table(doc, ['Escenario', 'Cantidad'], scen_data)

    doc.add_paragraph('')
    doc.add_heading('5.3 Días de Incapacidad', level=2)
    inj_data = [[i, str(c)] for i, c in sorted(by_injury.items(), key=lambda x: -x[1])]
    _add_table(doc, ['Días de Incapacidad', 'Cantidad'], inj_data)

    doc.add_page_break()

    # --- CONCLUSIONES ---
    h = doc.add_heading('6. Conclusiones y Recomendaciones', level=1)

    ai_conclusion = _ai_enhance(
        f"Basado en el análisis de {total} registros de lesiones no fatales en Colombia:\n"
        f"- Géneros: {dict(by_gender)}\n"
        f"- Contextos principales: {dict(by_context)}\n"
        f"- Departamentos con más casos: {dict(list(by_dept.items())[:5])}\n"
        f"- Días de incapacidad más comunes: {dict(by_injury)}\n\n"
        f"Genera 3 conclusiones clave y 3 recomendaciones accionables.",
        system_prompt="Eres un asesor en políticas públicas de seguridad ciudadana en Colombia."
    )

    if ai_conclusion:
        doc.add_paragraph(ai_conclusion)
    else:
        doc.add_paragraph(
            f"El análisis de {total} registros revela patrones importantes en la distribución "
            f"de lesiones no fatales. Se recomienda profundizar en los contextos de violencia "
            f"más frecuentes y fortalecer las estrategias de prevención en las zonas de mayor incidencia."
        )

    # --- FIRMA ---
    doc.add_paragraph('')
    doc.add_paragraph('')
    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run(f'\n\nGenerado automáticamente el {datetime.now().strftime("%d/%m/%Y %H:%M")}')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(150, 150, 150)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
